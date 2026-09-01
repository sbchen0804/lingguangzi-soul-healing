from pathlib import Path
from tempfile import TemporaryDirectory
import subprocess
import sys
import unittest
from unittest.mock import patch

from docx import Document
from PIL import Image
from reportlab.pdfgen.canvas import Canvas


SCRIPTS = Path(__file__).parents[2] / "skills/lingguangzi-chapter-site/scripts"
sys.path.insert(0, str(SCRIPTS))

from extract_lyrics import extract_lyrics
import render_pdf_pages as render_pdf_pages_module
from render_pdf_pages import _resolve_pdftoppm, render_pdf_pages


class LyricsTests(unittest.TestCase):
    def test_txt_lyrics_are_escaped_and_line_endings_are_normalized(self):
        with TemporaryDirectory() as raw:
            source = Path(raw) / "lyrics.txt"
            source.write_bytes("主歌 & 一\r\n風 < 心\r\n\r\n\r\n副歌\r\n今天".encode("utf-8"))

            result = extract_lyrics(source)

            self.assertEqual(set(result), {"status", "paragraphs", "html", "warnings"})
            self.assertEqual(result["status"], "ok")
            self.assertEqual(result["paragraphs"], ["主歌 & 一\n風 < 心", "副歌\n今天"])
            self.assertEqual(
                result["html"],
                "<p>主歌 &amp; 一<br>風 &lt; 心</p>\n<p>副歌<br>今天</p>",
            )
            self.assertEqual(result["warnings"], [])

    def test_docx_paragraphs_and_manual_line_breaks_are_preserved(self):
        with TemporaryDirectory() as raw:
            source = Path(raw) / "lyrics.docx"
            document = Document()
            first = document.add_paragraph()
            first.add_run("第一行 &")
            first.add_run().add_break()
            first.add_run("<第二行>")
            document.add_paragraph("副歌")
            document.save(source)

            result = extract_lyrics(source)

            self.assertEqual(result["status"], "ok")
            self.assertEqual(result["paragraphs"], ["第一行 &\n<第二行>", "副歌"])
            self.assertEqual(
                result["html"],
                "<p>第一行 &amp;<br>&lt;第二行&gt;</p>\n<p>副歌</p>",
            )

    def test_searchable_pdf_with_twenty_non_whitespace_characters_is_extracted(self):
        with TemporaryDirectory() as raw:
            source = Path(raw) / "lyrics.pdf"
            canvas = Canvas(str(source))
            canvas.drawString(72, 720, "12345678901234567890")
            canvas.save()

            result = extract_lyrics(source)

            self.assertEqual(result["status"], "ok")
            self.assertEqual(result["paragraphs"], ["12345678901234567890"])
            self.assertEqual(result["html"], "<p>12345678901234567890</p>")
            self.assertEqual(result["warnings"], [])

    def test_pdf_with_less_than_twenty_non_whitespace_characters_requests_ocr(self):
        with TemporaryDirectory() as raw:
            source = Path(raw) / "short.pdf"
            canvas = Canvas(str(source))
            canvas.drawString(72, 720, "1234567890123456789")
            canvas.save()

            result = extract_lyrics(source)

            self.assertEqual(result["status"], "requires_ocr")
            self.assertEqual(result["paragraphs"], [])
            self.assertEqual(result["html"], "")
            self.assertEqual(
                result["warnings"],
                ["PDF text extraction found fewer than 20 non-whitespace characters; OCR review is required."],
            )

    def test_scanned_pdf_requests_ocr_without_inventing_text(self):
        source = Path(__file__).parent / "fixtures/scanned-one-page.pdf"

        result = extract_lyrics(source)

        self.assertEqual(result["status"], "requires_ocr")
        self.assertEqual(result["paragraphs"], [])
        self.assertEqual(result["html"], "")
        self.assertEqual(
            result["warnings"],
            ["PDF text extraction found fewer than 20 non-whitespace characters; OCR review is required."],
        )


class PdfRenderTests(unittest.TestCase):
    def setUp(self):
        self.pdf = Path(__file__).parent / "fixtures/chapter-999/原圖文/original.pdf"

    def test_rendered_pages_are_ordered_named_and_have_dimensions(self):
        with TemporaryDirectory() as raw:
            source = Path(raw) / "two-pages.pdf"
            canvas = Canvas(str(source))
            canvas.drawString(72, 720, "page one")
            canvas.showPage()
            canvas.drawString(72, 720, "page two")
            canvas.save()
            output = Path(raw) / "pages"

            pages = render_pdf_pages(source, output, "original-999")

            self.assertEqual([page["page"] for page in pages], [1, 2])
            self.assertEqual(
                [Path(page["path"]).name for page in pages],
                ["original-999-page-01.png", "original-999-page-02.png"],
            )
            self.assertTrue(all(page["width"] >= 100 for page in pages))
            self.assertTrue(all(page["height"] >= 100 for page in pages))
            self.assertTrue(all(Path(page["path"]).is_file() for page in pages))

    def test_pdftoppm_resolution_prefers_path(self):
        path_program = Path(r"C:\tools\pdftoppm.exe")
        with patch.object(render_pdf_pages_module.shutil, "which", return_value=str(path_program)):
            self.assertEqual(_resolve_pdftoppm(), path_program)

    def test_pdftoppm_resolution_falls_back_to_bundled_runtime(self):
        with patch.object(render_pdf_pages_module.shutil, "which", return_value=None):
            resolved = _resolve_pdftoppm()
        self.assertTrue(resolved.is_file())
        self.assertEqual(resolved.name, "pdftoppm.exe")
        self.assertEqual(resolved.parent.name, "bin")

    def test_renderer_invokes_pdftoppm_with_an_argument_list(self):
        def create_page(args, **kwargs):
            Image.new("RGB", (200, 200), "white").save(Path(f"{args[-1]}-1.png"))
            return subprocess.CompletedProcess(args, 0, "", "")

        with TemporaryDirectory() as raw:
            with patch.object(render_pdf_pages_module.subprocess, "run", side_effect=create_page) as run:
                render_pdf_pages(self.pdf, Path(raw), "safe;prefix")

        command = run.call_args.args[0]
        options = run.call_args.kwargs
        self.assertIsInstance(command, list)
        self.assertEqual(command[1:8], ["-png", "-r", "144", "-f", "1", "-l", "1"])
        self.assertFalse(options.get("shell", False))

    def test_renderer_rejects_a_missing_page(self):
        completed = subprocess.CompletedProcess([], 0, "", "")
        with TemporaryDirectory() as raw:
            with patch.object(render_pdf_pages_module.subprocess, "run", return_value=completed):
                with self.assertRaisesRegex(RuntimeError, "missing rendered PDF page 1"):
                    render_pdf_pages(self.pdf, Path(raw), "original-999")

    def test_renderer_rejects_an_empty_page(self):
        def create_empty_page(args, **kwargs):
            Path(f"{args[-1]}-1.png").touch()
            return subprocess.CompletedProcess(args, 0, "", "")

        with TemporaryDirectory() as raw:
            with patch.object(render_pdf_pages_module.subprocess, "run", side_effect=create_empty_page):
                with self.assertRaisesRegex(RuntimeError, "empty rendered PDF page 1"):
                    render_pdf_pages(self.pdf, Path(raw), "original-999")

    def test_renderer_rejects_a_page_smaller_than_one_hundred_pixels(self):
        def create_small_page(args, **kwargs):
            Image.new("RGB", (99, 100), "white").save(Path(f"{args[-1]}-1.png"))
            return subprocess.CompletedProcess(args, 0, "", "")

        with TemporaryDirectory() as raw:
            with patch.object(render_pdf_pages_module.subprocess, "run", side_effect=create_small_page):
                with self.assertRaisesRegex(RuntimeError, "smaller than 100x100"):
                    render_pdf_pages(self.pdf, Path(raw), "original-999")


if __name__ == "__main__":
    unittest.main()
